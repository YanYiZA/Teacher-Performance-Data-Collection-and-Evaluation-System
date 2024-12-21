import sys
import pandas as pd
from docx import Document

def export_to_excel(data, file_name="output.xlsx"):
    """
    �����ݵ����� Excel �ļ�
    """
    # ����һ�� DataFrame
    df = pd.DataFrame(data[1:], columns=data[0])
    
    # ������ Excel
    df.to_excel(file_name, index=False, sheet_name="��Ч����")
    print(f"�����ѵ����� Excel �ļ���{file_name}")

def export_to_word(data, file_name="output.docx"):
    """
    �����ݵ����� Word �ļ�
    """
    # ����һ�� Word �ĵ�����
    doc = Document()
    
    # ��ӱ���
    doc.add_heading('��ʦ��Ч����', 0)
    
    # �������
    table = doc.add_table(rows=1, cols=len(data[0]))
    
    # ��ӱ�ͷ
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(data[0]):
        hdr_cells[i].text = col_name
    
    # �������
    for row in data[1:]:
        row_cells = table.add_row().cells
        for i, cell_value in enumerate(row):
            row_cells[i].text = str(cell_value)
    
    # �����ĵ�
    doc.save(file_name)
    print(f"�����ѵ����� Word �ļ���{file_name}")

def parse_data_from_stdin():
    """
    �ӱ�׼�����ȡ����
    """
    data = []
    
    for line in sys.stdin:
        # ����ÿһ�����ݣ�ȥ����β�Ļ��з����ö��ŷָ�
        row = line.strip().split(',')
        data.append(row)
    
    return data

def main():
    # ��ȡ�����в�������ʽ��
    format_type = sys.argv[1] if len(sys.argv) > 1 else "excel"

    # �ӱ�׼�����ȡ����
    data = parse_data_from_stdin()
    
    # ���ݸ�ʽѡ�񵼳�����
    if format_type == "excel":
        export_to_excel(data)
    elif format_type == "word":
        export_to_word(data)
    else:
        print("δ֪���ļ���ʽ")

if __name__ == "__main__":
    main()
